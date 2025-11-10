"""Helpers for deterministic simulation runs used by the mock audit flow."""
from __future__ import annotations

import hashlib
import json
import queue
import random
import shutil
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from zipfile import ZipFile, ZIP_DEFLATED

from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

try:
    from docx import Document
except ImportError:  # pragma: no cover - dependency optional in some environments
    Document = None  # type: ignore

ROOT = Path(__file__).resolve().parents[1]
MOCK_SIM_ROOT = ROOT / "mock_data" / "sim"
OUT_DIR = ROOT / "out"
AUDIT_LOG_DIR = ROOT / "audit_logs"

SIM_STREAMS: Dict[str, "queue.Queue[Optional[Dict[str, Any]]]"] = {}

DOCUMENT_TYPES = ("invoice", "purchase_order", "goods_receipt", "payment_advice")
ANOMALY_TYPES = (
    ("missing_po", "Missing purchase order linkage"),
    ("overpayment", "Payment exceeds invoice total"),
    ("duplicate_invoice", "Duplicate invoice detected"),
    ("currency_mismatch", "Currency mismatch between documents"),
)


@dataclass
class SyntheticDocument:
    doc_id: str
    doc_type: str
    vendor_id: str
    amount: float
    currency: str
    path: Path


def new_sim_run_id() -> str:
    """Return a timestamp-based identifier for simulation runs."""
    return datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S%f")


def _rand_for(run_id: str) -> random.Random:
    seed = hashlib.sha256(f"sim::{run_id}".encode("utf-8")).hexdigest()
    return random.Random(int(seed[:16], 16))


def _timestamp() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def get_stream_queue(run_id: str) -> "queue.Queue[Optional[Dict[str, Any]]]":
    return SIM_STREAMS.setdefault(run_id, queue.Queue())


def clear_stream_queue(run_id: str) -> None:
    SIM_STREAMS.pop(run_id, None)


def emit_sse_event(run_id: str, event_type: str, payload: Dict[str, Any]) -> None:
    queue_obj = get_stream_queue(run_id)
    queue_obj.put({"event": event_type, "payload": payload, "timestamp": _timestamp()})


def _ensure_run_dir(run_id: str) -> Path:
    run_dir = MOCK_SIM_ROOT / run_id
    run_dir.mkdir(parents=True, exist_ok=True)
    (run_dir / "files").mkdir(parents=True, exist_ok=True)
    return run_dir


def generate_synthetic_docs(run_id: str, vendor_id: str, sample_size: int = 20) -> List[SyntheticDocument]:
    rnd = _rand_for(run_id)
    run_dir = _ensure_run_dir(run_id)
    files_dir = run_dir / "files"

    documents: List[SyntheticDocument] = []
    for index in range(sample_size):
        doc_type = DOCUMENT_TYPES[index % len(DOCUMENT_TYPES)]
        doc_id = f"SIM-{run_id}-{index:03d}"
        currency = rnd.choice(["USD", "EUR", "GBP"])
        amount = round(rnd.uniform(500.0, 7500.0), 2)
        filename = f"{doc_id}-{doc_type}.txt"
        path = files_dir / filename
        path.write_text(
            json.dumps(
                {
                    "doc_id": doc_id,
                    "doc_type": doc_type,
                    "vendor_id": vendor_id,
                    "amount": amount,
                    "currency": currency,
                    "generated_at": _timestamp(),
                },
                indent=2,
            ),
            encoding="utf-8",
        )
        documents.append(
            SyntheticDocument(
                doc_id=doc_id,
                doc_type=doc_type,
                vendor_id=vendor_id,
                amount=amount,
                currency=currency,
                path=path,
            )
        )
    return documents


def inject_anomalies(documents: Sequence[SyntheticDocument], anomaly_rate: float, run_id: str) -> List[Dict[str, Any]]:
    if not documents:
        return []
    rnd = _rand_for(f"anomaly::{run_id}")
    threshold = max(1, int(len(documents) * anomaly_rate))
    selected = rnd.sample(list(range(len(documents))), k=threshold)

    anomalies: List[Dict[str, Any]] = []
    for counter, index in enumerate(sorted(selected)):
        doc = documents[index]
        anomaly_code, rationale = ANOMALY_TYPES[counter % len(ANOMALY_TYPES)]
        anomalies.append(
            {
                "id": f"SIM-ANOM-{run_id}-{counter:03d}",
                "document": doc.doc_id,
                "vendor": doc.vendor_id,
                "label": anomaly_code.replace("_", " ").title(),
                "severity": "high" if anomaly_code in {"overpayment", "duplicate_invoice"} else "medium",
                "rationale": rationale,
                "detected_at": _timestamp(),
            }
        )
    return anomalies


def seed_chat_timeline(run_id: str) -> List[Dict[str, Any]]:
    rnd = _rand_for(f"chat::{run_id}")
    prompts = [
        "Run vendor risk audit for the last quarter",
        "Highlight any late payments or missing purchase orders",
        "Summarize top anomalies affecting cashflow",
    ]
    responses = [
        "Starting synthetic audit pipeline and loading vendor data.",
        "Flagged a handful of invoices lacking PO references.",
        "Compiled summary of reconciliation progress for presentation.",
    ]
    keywords = [
        ["vendor risk", "quarterly"],
        ["late payment", "missing po"],
        ["summary", "cashflow"],
    ]
    base_timestamp = datetime(2025, 1, 1, 12, 0, tzinfo=timezone.utc)
    history: List[Dict[str, Any]] = []
    for idx, prompt in enumerate(prompts):
        ts_user = base_timestamp.isoformat().replace("+00:00", "Z")
        history.append(
            {
                "id": f"{run_id}-user-{idx}",
                "role": "user",
                "text": prompt,
                "timestamp": ts_user,
            }
        )
        offset_seconds = rnd.randint(15, 90)
        ts_assistant = (base_timestamp.replace(microsecond=0) + timedelta(seconds=offset_seconds)).isoformat().replace("+00:00", "Z")
        history.append(
            {
                "id": f"{run_id}-assistant-{idx}",
                "role": "assistant",
                "text": responses[idx],
                "timestamp": ts_assistant,
                "keywords": keywords[idx],
            }
        )
        base_timestamp += timedelta(minutes=2)
    return history


def mock_policy_check(run_context: Dict[str, Any], policy_file: Optional[Path] = None) -> List[Dict[str, Any]]:
    vendor_id = run_context.get("vendor_id", "VEND-SIM")
    violations: List[Dict[str, Any]] = []
    documents = run_context.get("documents", [])
    doc_ids = [doc.doc_id if isinstance(doc, SyntheticDocument) else doc.get("doc_id") for doc in documents]
    if not doc_ids:
        doc_ids = ["SIM-DOC-001"]
    controls = [
        ("SOX_404", "Dual approval missing for high-value invoice."),
        ("VENDOR_RISK", "Vendor risk scoring not documented for new suppliers."),
    ]
    for index, (control_id, statement) in enumerate(controls):
        violations.append(
            {
                "id": f"SIM-POL-{run_context.get('run_id', 'demo')}-{index:02d}",
                "control": control_id,
                "control_label": control_id.replace("_", " "),
                "statement": statement,
                "evidence_excerpt": f"Related document {doc_ids[index % len(doc_ids)]} for vendor {vendor_id}.",
                "severity": "high" if index == 0 else "medium",
                "confidence": 0.9 if index == 0 else 0.75,
                "page": index + 1,
            }
        )
    return violations


def build_manifest_and_package(run_context: Dict[str, Any]) -> Path:
    run_id = run_context["run_id"]
    run_dir = _ensure_run_dir(run_id)
    manifest_path = run_dir / "manifest.json"
    summary_path = run_dir / "summary.md"

    manifest = {
        "run_id": run_id,
        "vendor_id": run_context.get("vendor_id"),
        "generated_at": _timestamp(),
        "documents": [
            {
                "doc_id": doc.doc_id,
                "doc_type": doc.doc_type,
                "amount": doc.amount,
                "currency": doc.currency,
                "path": str(doc.path.relative_to(ROOT)),
            }
            for doc in run_context.get("documents", [])
            if isinstance(doc, SyntheticDocument)
        ],
        "anomalies": run_context.get("anomalies", []),
        "policy_violations": run_context.get("policy_violations", []),
    }
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    summary_lines = [
        f"# Audit Simulation Report {run_id}",
        "",
        f"Vendor **{run_context.get('vendor_id')}**",
        f"Generated at {_timestamp()}",
        "",
        "## Highlights",
        f"- Documents processed: {len(run_context.get('documents', []))}",
        f"- Anomalies detected: {len(run_context.get('anomalies', []))}",
        f"- Policy violations: {len(run_context.get('policy_violations', []))}",
    ]
    summary_path.write_text("\n".join(summary_lines), encoding="utf-8")

    package_path = OUT_DIR / f"package_SIM_{run_id}.zip"
    package_path.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(package_path, "w", ZIP_DEFLATED) as archive:
        archive.write(manifest_path, arcname=f"SIM_{run_id}/manifest.json")
        archive.write(summary_path, arcname=f"SIM_{run_id}/summary.md")
        for doc in run_context.get("documents", []):
            if isinstance(doc, SyntheticDocument):
                archive.write(doc.path, arcname=f"SIM_{run_id}/files/{doc.path.name}")
    run_context["manifest_path"] = manifest_path
    run_context["package_path"] = package_path
    return manifest_path


def generate_pdf_docx_report(run_context: Dict[str, Any]) -> Tuple[Path, Optional[Path]]:
    run_id = run_context["run_id"]
    reports_dir = OUT_DIR / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = reports_dir / f"SIM_{run_id}.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=LETTER)
    width, height = LETTER
    text_obj = c.beginText(72, height - 72)
    text_obj.setFont("Helvetica", 12)
    lines = [
        f"AudItecX Simulation Report {run_id}",
        "",
        f"Vendor: {run_context.get('vendor_id')}",
        f"Documents processed: {len(run_context.get('documents', []))}",
        f"Anomalies detected: {len(run_context.get('anomalies', []))}",
        f"Policy violations: {len(run_context.get('policy_violations', []))}",
        f"Generated at: {_timestamp()}",
    ]
    for line in lines:
        text_obj.textLine(line)
    c.drawText(text_obj)
    c.showPage()
    c.save()

    docx_path: Optional[Path] = None
    if Document is not None:
        doc = Document()
        doc.add_heading("AudItecX Simulation Report", level=0)
        doc.add_paragraph(f"Run ID: {run_id}")
        doc.add_paragraph(f"Vendor: {run_context.get('vendor_id')}")
        doc.add_paragraph(f"Documents processed: {len(run_context.get('documents', []))}")
        doc.add_paragraph(f"Anomalies detected: {len(run_context.get('anomalies', []))}")
        doc.add_paragraph(f"Policy violations: {len(run_context.get('policy_violations', []))}")
        doc.add_paragraph(f"Generated at: {_timestamp()}")
        docx_path = reports_dir / f"SIM_{run_id}.docx"
        doc.save(docx_path)

    run_context["report_pdf_path"] = pdf_path
    run_context["report_docx_path"] = docx_path
    return pdf_path, docx_path


def compare_with_last_run(current_manifest: Dict[str, Any], previous_manifest: Optional[Dict[str, Any]]) -> str:
    if not previous_manifest:
        return "No prior runs available for comparison."
    current_docs = {doc["doc_id"] for doc in current_manifest.get("documents", [])}
    previous_docs = {doc["doc_id"] for doc in previous_manifest.get("documents", [])}
    added = current_docs - previous_docs
    removed = previous_docs - current_docs
    fragments = []
    if added:
        fragments.append(f"Added {len(added)} documents")
    if removed:
        fragments.append(f"Removed {len(removed)} documents")
    if not fragments:
        fragments.append("Document set unchanged from previous run")
    return "; ".join(fragments)


def cleanup_simulation_artifacts(run_id: str) -> None:
    run_dir = MOCK_SIM_ROOT / run_id
    if run_dir.exists():
        shutil.rmtree(run_dir)
    package_path = OUT_DIR / f"package_SIM_{run_id}.zip"
    if package_path.exists():
        package_path.unlink()
    reports_dir = OUT_DIR / "reports"
    for ext in (".pdf", ".docx"):
        candidate = reports_dir / f"SIM_{run_id}{ext}"
        if candidate.exists():
            candidate.unlink()
    audit_log = AUDIT_LOG_DIR / f"SIM_{run_id}.json"
    if audit_log.exists():
        audit_log.unlink()
    clear_stream_queue(run_id)