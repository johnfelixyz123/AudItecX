"""Flask UI for executing AudItecX orchestrations."""
from __future__ import annotations

import calendar
import json
import logging
import queue
import threading
import time
import re
import uuid
import textwrap
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from flask import Response, Flask, jsonify, render_template, request, send_file, stream_with_context
from flask.typing import ResponseReturnValue
from werkzeug.utils import secure_filename

import sys

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "src" / "ui" / "templates"
if str(ROOT) not in sys.path:
    sys.path.append(str(ROOT))

from src.agents import planner_agent  # noqa: E402
from src.nlu import intent_parser  # noqa: E402
from src.orchestrator import Orchestrator  # noqa: E402
from src.packager import AUDIT_LOG_DIR, OUT_DIR  # noqa: E402
from src.policy_checker import DEFAULT_CONTROL_IDS, analyze_policy_document  # noqa: E402
from src.vendor_metrics import VendorRiskError, build_vendor_risk_metrics  # noqa: E402


app = Flask(__name__, template_folder=str(TEMPLATE_DIR))
app.config["SECRET_KEY"] = "auditecx-dev"

logging.basicConfig(level=logging.INFO)
LOGGER = logging.getLogger("auditecx.app")

orchestrator = Orchestrator()
RUN_QUEUES: Dict[str, "queue.Queue[Dict[str, Any] | None]"] = {}
RUN_RESULTS: Dict[str, Dict[str, Any]] = {}
RUN_LOCK = threading.Lock()
NOTIFICATION_LOCK = threading.Lock()
SCHEDULE_LOCK = threading.Lock()
SCHEDULER_STOP = threading.Event()

SCHEDULE_FREQUENCIES = {"daily", "weekly", "monthly"}
SCHEDULER_POLL_SECONDS = 60


def _new_run_id() -> str:
    return datetime.utcnow().strftime("%Y%m%d%H%M%S%f")


def _enqueue(run_id: str, event: str, payload: Dict[str, Any]) -> None:
    with RUN_LOCK:
        q = RUN_QUEUES.get(run_id)
    if q:
        q.put({"event": event, "payload": payload})


def _conversation_path(run_id: str) -> Path:
    AUDIT_LOG_DIR.mkdir(parents=True, exist_ok=True)
    return AUDIT_LOG_DIR / f"{run_id}_conversation.json"


def _timestamp() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


VENDOR_TOKEN_PATTERN = re.compile(r"VEND-[A-Za-z0-9-]+")
HEATMAP_TIMESTAMP_KEYS = ("timestamp", "created_at", "detected_at", "occurred_at")


def _extract_keywords(text: str, limit: int = 8) -> List[str]:
    tokens = re.findall(r"[A-Za-z0-9]{4,}", text.lower())
    if not tokens:
        return []
    counts = Counter(tokens)
    return [word for word, _ in counts.most_common(limit)]


def _build_message(run_id: str, role: str, text: str, index: int) -> Dict[str, Any]:
    message: Dict[str, Any] = {
        "id": f"{run_id}-{role}-{index}",
        "role": role,
        "text": text,
        "timestamp": _timestamp(),
    }
    keywords = _extract_keywords(text)
    if keywords:
        message["keywords"] = keywords
    return message


def _persist_conversation(run_id: str, messages: List[Dict[str, Any]]) -> None:
    path = _conversation_path(run_id)
    payload = {
        "run_id": run_id,
        "messages": messages,
        "updated_at": _timestamp(),
    }
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _notifications_path() -> Path:
    AUDIT_LOG_DIR.mkdir(parents=True, exist_ok=True)
    return AUDIT_LOG_DIR / "notifications.json"


def _schedules_path() -> Path:
    AUDIT_LOG_DIR.mkdir(parents=True, exist_ok=True)
    return AUDIT_LOG_DIR / "schedules.json"


def _load_notifications() -> List[Dict[str, Any]]:
    path = _notifications_path()
    if not path.exists():
        return []
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        LOGGER.warning("Failed to parse notifications log at %s", path)
        return []

    if isinstance(raw, dict) and isinstance(raw.get("notifications"), list):
        payload = raw.get("notifications", [])
    else:
        payload = raw

    if not isinstance(payload, list):
        return []

    notifications: List[Dict[str, Any]] = []
    for entry in payload:
        if isinstance(entry, dict):
            notifications.append(entry)
    return notifications


def _write_notifications(notifications: Iterable[Dict[str, Any]]) -> None:
    path = _notifications_path()
    data = list(notifications)
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def _append_notification(notification: Dict[str, Any]) -> None:
    with NOTIFICATION_LOCK:
        existing = _load_notifications()
        existing.append(notification)
        _write_notifications(existing)


def _build_notification(notification_type: str, message: str) -> Dict[str, Any]:
    return {
        "id": str(uuid.uuid4()),
        "type": notification_type,
        "message": message,
        "timestamp": _timestamp(),
        "read": False,
    }


def _record_run_notifications(run_id: str, result: Dict[str, Any]) -> None:
    _append_notification(_build_notification("run_complete", f"Audit run {run_id} completed."))

    context = result.get("context") if isinstance(result, dict) else None
    anomalies = context.get("anomalies") if isinstance(context, dict) else None
    if isinstance(anomalies, list) and anomalies:
        count = len(anomalies)
        noun = "anomaly" if count == 1 else "anomalies"
        _append_notification(
            _build_notification(
                "anomaly_alert",
                f"{count} {noun} pending review for run {run_id}.",
            )
        )


def _reports_dir() -> Path:
    path = OUT_DIR / "reports"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _load_schedules() -> List[Dict[str, Any]]:
    path = _schedules_path()
    if not path.exists():
        return []
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        LOGGER.warning("Failed to parse schedules log at %s", path)
        return []

    if isinstance(raw, dict) and isinstance(raw.get("schedules"), list):
        payload = raw.get("schedules", [])
    else:
        payload = raw

    if not isinstance(payload, list):
        return []

    schedules: List[Dict[str, Any]] = []
    for entry in payload:
        if isinstance(entry, dict):
            schedules.append(entry)
    return schedules


def _write_schedules(schedules: Iterable[Dict[str, Any]]) -> None:
    path = _schedules_path()
    data = list(schedules)
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def _isoformat_utc(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")


def _parse_datetime(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    cleaned = value.strip()
    if not cleaned:
        return None
    cleaned = cleaned.replace("Z", "+00:00")
    try:
        dt = datetime.fromisoformat(cleaned)
    except ValueError:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc)


def _parse_start_date(value: Optional[str]) -> datetime:
    if not value:
        return datetime.now(timezone.utc)
    cleaned = value.strip()
    if not cleaned:
        return datetime.now(timezone.utc)
    cleaned = cleaned.replace("Z", "")
    for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M", "%Y-%m-%dT%H:%M:%S"):
        try:
            dt = datetime.strptime(cleaned, fmt)
            dt = dt.replace(tzinfo=timezone.utc)
            return dt
        except ValueError:
            continue
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return datetime.now(timezone.utc)
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc)


def _advance_month(dt: datetime) -> datetime:
    year = dt.year + (dt.month // 12)
    month = dt.month % 12 + 1
    day = min(dt.day, calendar.monthrange(year, month)[1])
    return dt.replace(year=year, month=month, day=day)


def _increment_schedule_time(dt: datetime, frequency: str) -> Optional[datetime]:
    if frequency == "daily":
        return dt + timedelta(days=1)
    if frequency == "weekly":
        return dt + timedelta(weeks=1)
    if frequency == "monthly":
        return _advance_month(dt)
    return None


def _determine_next_run(start: datetime, frequency: str, now: Optional[datetime] = None) -> Optional[datetime]:
    if frequency not in SCHEDULE_FREQUENCIES:
        return None
    current = start.astimezone(timezone.utc)
    reference = now or datetime.now(timezone.utc)
    if current > reference:
        return current
    for _ in range(512):
        next_candidate = _increment_schedule_time(current, frequency)
        if next_candidate is None:
            return None
        current = next_candidate
        if current > reference:
            return current
    return None


def _trigger_scheduled_run(schedule: Dict[str, Any]) -> str:
    vendor_id = str(schedule.get("vendor_id") or "Unknown Vendor").strip() or "Unknown Vendor"
    text = f"Run a scheduled audit for vendor {vendor_id}."
    parsed_intent = intent_parser.parse_intent(text, use_llm=False)
    plan = planner_agent.plan_tasks(parsed_intent)

    run_id = _new_run_id()
    thread = threading.Thread(target=_worker, args=(run_id, parsed_intent, plan, None, text), daemon=True)
    thread.start()

    message = f"Scheduled audit triggered for {vendor_id} (run {run_id})."
    _append_notification(_build_notification("scheduler_triggered", message))
    return run_id


def _poll_schedules() -> None:
    now = datetime.now(timezone.utc)
    updated = False
    with SCHEDULE_LOCK:
        schedules = _load_schedules()
        for schedule in schedules:
            frequency = str(schedule.get("frequency") or "").lower()
            next_run = _parse_datetime(schedule.get("next_run_at"))
            if frequency not in SCHEDULE_FREQUENCIES or next_run is None:
                continue
            if next_run <= now:
                run_id = _trigger_scheduled_run(schedule)
                schedule["last_run_at"] = _isoformat_utc(now)
                schedule["last_run_id"] = run_id
                future = _determine_next_run(next_run, frequency, now)
                schedule["next_run_at"] = _isoformat_utc(future) if future else None
                updated = True
        if updated:
            _write_schedules(schedules)


def _scheduler_loop() -> None:
    while not SCHEDULER_STOP.is_set():
        try:
            _poll_schedules()
        except Exception:  # pragma: no cover - defensive logging for background loop
            LOGGER.exception("Scheduler loop failed")
        SCHEDULER_STOP.wait(SCHEDULER_POLL_SECONDS)


SCHEDULER_THREAD = threading.Thread(target=_scheduler_loop, daemon=True)
SCHEDULER_THREAD.start()


def _load_run_manifest(run_id: str) -> Optional[Dict[str, Any]]:
    manifest_path = AUDIT_LOG_DIR / f"{run_id}.json"
    if not manifest_path.exists():
        return None
    try:
        return json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        LOGGER.warning("Failed to read manifest for run %s", run_id)
        return None


def _markdown_to_plain_text(markdown_value: str) -> str:
    cleaned = re.sub(r"`{1,3}.*?`{1,3}", "", markdown_value, flags=re.DOTALL)
    cleaned = re.sub(r"[#>*_]+", "", cleaned)
    lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
    return " ".join(lines)


def _resolve_run_snapshot(run_id: str) -> Optional[Dict[str, Any]]:
    result = RUN_RESULTS.get(run_id)
    manifest = _load_run_manifest(run_id)

    if not result and not manifest:
        return None

    snapshot: Dict[str, Any] = {
        "run_id": run_id,
        "context": {},
        "summary_text": None,
        "summary_markdown": None,
        "anomalies": [],
        "documents": [],
        "journal_entries": [],
        "totals": {},
    }

    if isinstance(result, dict):
        context = result.get("context")
        if isinstance(context, dict):
            snapshot["context"] = context
            summary_text = context.get("summary_text") or context.get("summary")
            if isinstance(summary_text, str):
                snapshot["summary_text"] = summary_text
            anomalies = context.get("anomalies")
            if isinstance(anomalies, list):
                snapshot["anomalies"] = anomalies
            documents = context.get("documents")
            if isinstance(documents, list):
                snapshot["documents"] = documents
            journal_entries = context.get("journal_entries")
            if isinstance(journal_entries, list):
                snapshot["journal_entries"] = journal_entries
            totals = context.get("totals")
            if isinstance(totals, dict):
                snapshot["totals"] = totals

        summary_text = result.get("summary_text")
        if isinstance(summary_text, str):
            snapshot["summary_text"] = summary_text

    if manifest:
        summary_block = manifest.get("summary")
        if isinstance(summary_block, dict):
            markdown_value = summary_block.get("markdown")
            if isinstance(markdown_value, str):
                snapshot["summary_markdown"] = markdown_value
            anomalies = summary_block.get("anomalies")
            if isinstance(anomalies, list) and not snapshot["anomalies"]:
                snapshot["anomalies"] = anomalies
            totals = summary_block.get("totals")
            if isinstance(totals, dict) and not snapshot["totals"]:
                snapshot["totals"] = totals

        documents = manifest.get("documents")
        if isinstance(documents, list) and not snapshot["documents"]:
            snapshot["documents"] = documents
        journal_entries = manifest.get("journal_entries")
        if isinstance(journal_entries, list) and not snapshot["journal_entries"]:
            snapshot["journal_entries"] = journal_entries

    return snapshot


def _format_anomaly_entry(entry: Any) -> str:
    if isinstance(entry, str):
        return entry.strip()
    if isinstance(entry, dict):
        label = str(entry.get("label") or entry.get("title") or entry.get("id") or "Anomaly").strip()
        detail = entry.get("rationale") or entry.get("description") or entry.get("message") or entry.get("details")
        if isinstance(detail, str) and detail.strip():
            return f"{label}: {detail.strip()}"
        return label
    return str(entry)


def _format_document_entry(entry: Any) -> str:
    if not isinstance(entry, dict):
        return str(entry)
    name = str(entry.get("document") or entry.get("filename") or entry.get("name") or "Supporting document").strip()
    vendor = str(entry.get("vendor") or entry.get("vendor_name") or entry.get("vendor_id") or "N/A").strip()
    invoice = str(entry.get("invoice") or entry.get("invoice_id") or "").strip()
    currency = str(entry.get("currency") or "").strip()
    amount_value = entry.get("amount")
    amount = None
    if isinstance(amount_value, (int, float)):
        amount = f"{amount_value:,.2f}"
    elif isinstance(amount_value, str) and amount_value.strip():
        amount = amount_value.strip()
    fragments = [name, f"Vendor {vendor}"]
    if invoice:
        fragments.append(f"Invoice {invoice}")
    if amount:
        fragments.append(f"Amount {amount}{f' {currency}' if currency else ''}")
    return " — ".join(fragments)


def _compose_report_content(snapshot: Dict[str, Any]) -> Dict[str, Any]:
    generated_at = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")

    summary_text: Optional[str] = snapshot.get("summary_text") if isinstance(snapshot.get("summary_text"), str) else None
    if not summary_text and isinstance(snapshot.get("summary_markdown"), str):
        summary_text = _markdown_to_plain_text(snapshot["summary_markdown"])
    if not summary_text:
        summary_text = "Summary data unavailable for this run."

    raw_totals = snapshot.get("totals") if isinstance(snapshot.get("totals"), dict) else {}
    totals: Dict[str, str] = {}
    for key, value in raw_totals.items():
        label = str(key).replace("_", " ").title()
        if isinstance(value, (int, float)):
            totals[label] = f"{value:,.2f}"
        else:
            totals[label] = str(value)

    anomalies_raw = snapshot.get("anomalies") if isinstance(snapshot.get("anomalies"), list) else []
    anomalies = [_format_anomaly_entry(item) for item in anomalies_raw][:10]

    documents_raw = snapshot.get("documents") if isinstance(snapshot.get("documents"), list) else []
    documents = [_format_document_entry(item) for item in documents_raw][:6]

    journal_entries = snapshot.get("journal_entries") if isinstance(snapshot.get("journal_entries"), list) else []

    metrics = {
        "documents_total": len(documents_raw),
        "anomalies_total": len(anomalies_raw),
        "journal_entries_total": len(journal_entries),
    }

    lines: List[str] = [
        "AudItecX Audit Report",
        f"Run ID: {snapshot['run_id']}",
        f"Generated: {generated_at}",
        "",
        "Executive Summary",
    ]

    for wrapped in textwrap.wrap(summary_text, width=90):
        lines.append(wrapped)

    if totals:
        lines.extend(["", "Key Totals"])
        for label, value in totals.items():
            lines.append(f"- {label}: {value}")

    if anomalies:
        lines.extend(["", "Anomalies"])
        for entry in anomalies:
            lines.append(f"- {entry}")
    else:
        lines.extend(["", "Anomalies", "- None recorded"])

    if documents:
        lines.extend(["", "Evidence Snapshot"])
        for entry in documents:
            lines.append(f"- {entry}")

    lines.extend(["", "Generated by AudItecX prototype (mock mode)."])

    return {
        "run_id": snapshot["run_id"],
        "generated_at": generated_at,
        "summary": summary_text,
        "totals": totals,
        "anomalies": anomalies,
        "documents": documents,
        "metrics": metrics,
        "lines": lines,
    }


def _render_pdf(path: Path, content: Dict[str, Any]) -> None:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
    except ImportError as exc:  # pragma: no cover - dependency issue handled at runtime
        raise RuntimeError("PDF generation requires reportlab") from exc

    c = canvas.Canvas(str(path), pagesize=LETTER)
    width, height = LETTER
    text_object = c.beginText(0.8 * inch, height - 1 * inch)
    text_object.setFont("Helvetica", 11)

    for line in content.get("lines", []):
        if text_object.getY() <= 0.8 * inch:
            c.drawText(text_object)
            c.showPage()
            text_object = c.beginText(0.8 * inch, height - 1 * inch)
            text_object.setFont("Helvetica", 11)
        text_object.textLine(line)

    c.drawText(text_object)
    c.showPage()
    c.save()


def _render_docx(path: Path, content: Dict[str, Any]) -> None:
    try:
        from docx import Document
    except ImportError:  # pragma: no cover - dependency issue handled at runtime
        _render_docx_fallback(path, content)
        return

    document = Document()
    document.add_heading("AudItecX Audit Report", level=0)
    document.add_paragraph(f"Run ID: {content.get('run_id')}")
    document.add_paragraph(f"Generated: {content.get('generated_at')}")

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(content.get("summary", "Summary data unavailable for this run."))

    totals = content.get("totals") if isinstance(content.get("totals"), dict) else {}
    if totals:
        document.add_heading("Key Totals", level=1)
        table = document.add_table(rows=1, cols=2)
        header = table.rows[0].cells
        header[0].text = "Metric"
        header[1].text = "Value"
        for label, value in totals.items():
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

    anomalies = content.get("anomalies") if isinstance(content.get("anomalies"), list) else []
    document.add_heading("Anomalies", level=1)
    if anomalies:
        for entry in anomalies:
            document.add_paragraph(entry, style="List Bullet")
    else:
        document.add_paragraph("None recorded")

    documents = content.get("documents") if isinstance(content.get("documents"), list) else []
    document.add_heading("Evidence Snapshot", level=1)
    if documents:
        for entry in documents:
            document.add_paragraph(entry, style="List Number")
    else:
        document.add_paragraph("No supporting documents captured in this snapshot.")

    document.add_paragraph("")
    document.add_paragraph("Generated by AudItecX prototype (mock mode).")

    document.save(path)


def _render_docx_fallback(path: Path, content: Dict[str, Any]) -> None:
    """Create a minimal DOCX archive when python-docx is unavailable."""
    from xml.sax.saxutils import escape
    from zipfile import ZipFile, ZIP_DEFLATED

    lines: List[str] = content.get("lines") or []
    paragraphs: List[str] = []
    for line in lines:
        if not line:
            paragraphs.append("<w:p/>")
            continue
        paragraphs.append(
            "<w:p><w:r><w:t xml:space=\"preserve\">{}</w:t></w:r></w:p>".format(escape(line))
        )

    body = "".join(paragraphs) + (
        "<w:sectPr><w:pgSz w:w=\"12240\" w:h=\"15840\"/>"
        "<w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\""
        " w:header=\"720\" w:footer=\"720\" w:gutter=\"0\"/></w:sectPr>"
    )

    document_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        f"<w:body>{body}</w:body>"
        "</w:document>"
    )

    styles_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:styles xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        "<w:style w:type=\"paragraph\" w:default=\"1\" w:styleId=\"Normal\">"
        "<w:name w:val=\"Normal\"/>"
        "</w:style>"
        "</w:styles>"
    )

    rels_root = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
        "<Relationship Id=\"R1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\""
        " Target=\"word/document.xml\"/>"
        "</Relationships>"
    )

    document_rels = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\"/>"
    )

    content_types = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
        "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
        "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
        "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
        "<Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
        "</Types>"
    )

    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels_root)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", document_rels)
        archive.writestr("word/styles.xml", styles_xml)


def _extract_anomalies_from_payload(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    context = payload.get("context") if isinstance(payload, dict) else None
    candidates = []
    if isinstance(context, dict):
        anomalies = context.get("anomalies")
        if isinstance(anomalies, list):
            candidates = anomalies
    if not candidates and isinstance(payload, dict):
        anomalies = payload.get("anomalies")
        if isinstance(anomalies, list):
            candidates = anomalies
    return [entry for entry in candidates if isinstance(entry, dict)]


def _collect_anomaly_records() -> List[Tuple[Optional[str], Dict[str, Any]]]:
    records: List[Tuple[Optional[str], Dict[str, Any]]] = []

    if AUDIT_LOG_DIR.exists():
        for path in AUDIT_LOG_DIR.glob("*.json"):
            if path.name.endswith("_conversation.json") or path.name == "notifications.json":
                continue
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                LOGGER.warning("Failed to parse audit log %s", path)
                continue
            anomalies = _extract_anomalies_from_payload(data)
            if not anomalies:
                continue
            run_id = str(data.get("run_id") or data.get("context", {}).get("run_id") or path.stem)
            for anomaly in anomalies:
                records.append((run_id, anomaly))

    if not records and OUT_DIR.exists():
        for manifest_path in OUT_DIR.glob("*/manifest.json"):
            try:
                data = json.loads(manifest_path.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                LOGGER.warning("Failed to parse manifest %s", manifest_path)
                continue
            anomalies = _extract_anomalies_from_payload(data)
            if not anomalies:
                continue
            run_id = str(data.get("run_id") or manifest_path.parent.name)
            for anomaly in anomalies:
                records.append((run_id, anomaly))

    return records


def _extract_vendor_from_anomaly(anomaly: Dict[str, Any]) -> Optional[str]:
    vendor = anomaly.get("vendor_id") or anomaly.get("vendor")
    if isinstance(vendor, str) and vendor.strip():
        return vendor.strip()

    for key in ("ledger_entry", "document", "metadata", "details"):
        entry = anomaly.get(key)
        if isinstance(entry, dict):
            candidate = entry.get("vendor_id") or entry.get("vendor")
            if isinstance(candidate, str) and candidate.strip():
                return candidate.strip()

    doc_field = anomaly.get("doc")
    if isinstance(doc_field, str):
        match = VENDOR_TOKEN_PATTERN.search(doc_field)
        if match:
            return match.group(0)

    return None


def _parse_month_from_string(value: str) -> Optional[str]:
    if not value:
        return None
    cleaned = value.strip()
    if not cleaned:
        return None
    cleaned = cleaned.replace("Z", "+00:00")
    try:
        dt = datetime.fromisoformat(cleaned)
        return dt.strftime("%Y-%m")
    except ValueError:
        pass

    digits = re.findall(r"\d", cleaned)
    if digits:
        joined = "".join(digits)
        try:
            if len(joined) >= 8:
                dt = datetime.strptime(joined[:8], "%Y%m%d")
            else:
                dt = datetime.strptime(joined[:6], "%Y%m")
            return dt.strftime("%Y-%m")
        except ValueError:
            pass

    match = re.search(r"(20\d{2})[-/](\d{2})", cleaned)
    if match:
        return f"{match.group(1)}-{match.group(2)}"

    return None


def _resolve_run_month(run_id: Optional[str], anomaly: Dict[str, Any]) -> Optional[str]:
    candidate = None
    if run_id:
        digits = re.findall(r"\d", run_id)
        if digits:
            joined = "".join(digits)
            if len(joined) >= 8:
                try:
                    candidate = datetime.strptime(joined[:8], "%Y%m%d").strftime("%Y-%m")
                except ValueError:
                    candidate = None
            elif len(joined) >= 6:
                try:
                    candidate = datetime.strptime(joined[:6], "%Y%m").strftime("%Y-%m")
                except ValueError:
                    candidate = None
    if candidate:
        return candidate

    for key in HEATMAP_TIMESTAMP_KEYS:
        value = anomaly.get(key)
        if isinstance(value, str):
            month = _parse_month_from_string(value)
            if month:
                return month

    return None


def _build_heatmap_counts(mode: str) -> Dict[str, int]:
    records = _collect_anomaly_records()
    if not records:
        return {}

    counts: Dict[str, int] = defaultdict(int)

    if mode == "vendor":
        for _run_id, anomaly in records:
            vendor_id = _extract_vendor_from_anomaly(anomaly)
            if vendor_id:
                counts[vendor_id] += 1
    else:
        for run_id, anomaly in records:
            month = _resolve_run_month(run_id, anomaly)
            if month:
                counts[month] += 1

    return counts


def _format_heatmap_payload(mode: str) -> Dict[str, List[str] | List[int]]:
    counts = _build_heatmap_counts(mode)
    if not counts:
        return {"labels": [], "values": []}

    if mode == "vendor":
        items = sorted(counts.items(), key=lambda item: (-item[1], item[0]))
    else:
        items = sorted(counts.items(), key=lambda item: item[0])

    labels = [label for label, _value in items]
    values = [value for _label, value in items]
    return {"labels": labels, "values": values}


def _worker(run_id: str, parsed_intent: Dict[str, Any], plan: Any, email: str | None, source_text: str) -> None:
    def stream_callback(chunk: str) -> None:
        assistant_chunks.append(chunk)
        _ensure_assistant_entry("".join(assistant_chunks))
        _persist_conversation(run_id, conversation_messages)
        _enqueue(run_id, "summary_chunk", {"text": chunk})

    assistant_chunks: List[str] = []
    message_index = 0
    conversation_messages: List[Dict[str, Any]] = []
    assistant_entry: Dict[str, Any] | None = None

    if source_text:
        conversation_messages.append(_build_message(run_id, "user", source_text, message_index))
        message_index += 1
        _persist_conversation(run_id, conversation_messages)

    def _ensure_assistant_entry(text: str) -> None:
        nonlocal assistant_entry, message_index
        if assistant_entry is None:
            assistant_entry = _build_message(run_id, "assistant", text, message_index)
            message_index += 1
            conversation_messages.append(assistant_entry)
        else:
            assistant_entry["text"] = text
            keywords = _extract_keywords(text)
            if keywords:
                assistant_entry["keywords"] = keywords
            elif "keywords" in assistant_entry:
                assistant_entry.pop("keywords")
            assistant_entry["timestamp"] = _timestamp()

    try:
        _enqueue(run_id, "status", {"message": "Starting orchestration"})
        result = orchestrator.execute(
            parsed_intent=parsed_intent,
            plan=plan,
            stream_callback=stream_callback,
            run_id=run_id,
        )
        summary_text = result.get("context", {}).get("summary_text") or "".join(assistant_chunks)
        if summary_text:
            _ensure_assistant_entry(summary_text)
            _persist_conversation(run_id, conversation_messages)
        result["conversation"] = conversation_messages
        RUN_RESULTS[run_id] = result
        try:
            _record_run_notifications(run_id, result)
        except Exception:  # pragma: no cover - notification logging should not break runs
            LOGGER.exception("Failed to record notification for run %s", run_id)
        package_path = result.get("package_path")
        _enqueue(
            run_id,
            "complete",
            {
                "package_path": package_path,
                "manifest_path": result.get("manifest_path"),
                "email": email,
            },
        )
    except Exception as exc:  # pragma: no cover - streamed to client
        LOGGER.exception("Run %s failed", run_id)
        _enqueue(run_id, "error", {"message": str(exc)})
    finally:
        if conversation_messages:
            try:
                _persist_conversation(run_id, conversation_messages)
            except OSError:
                LOGGER.exception("Failed to persist conversation for %s", run_id)
        with RUN_LOCK:
            q = RUN_QUEUES.get(run_id)
        if q:
            q.put(None)
        with RUN_LOCK:
            RUN_QUEUES.pop(run_id, None)


@app.route("/")
def index() -> str:
    return render_template("index.html")


@app.get("/api/vendors/risk")
def vendor_risk() -> ResponseReturnValue:
    try:
        payload = build_vendor_risk_metrics()
    except VendorRiskError as exc:
        LOGGER.exception("Failed to build vendor risk metrics")
        return jsonify({"error": str(exc)}), 500
    return jsonify(payload)


@app.get("/api/vendors/heatmap")
def vendor_anomaly_heatmap() -> ResponseReturnValue:
    mode = (request.args.get("by") or "vendor").lower()
    if mode not in {"vendor", "month"}:
        return jsonify({"error": "Query parameter 'by' must be 'vendor' or 'month'"}), 400

    payload = _format_heatmap_payload(mode)
    return jsonify(payload)


@app.get("/api/notifications")
def notifications_list() -> ResponseReturnValue:
    notifications = sorted(_load_notifications(), key=lambda item: item.get("timestamp", ""), reverse=True)
    return jsonify({"notifications": notifications})


@app.post("/api/notifications/ack")
def notifications_ack() -> ResponseReturnValue:
    data = request.get_json(force=True, silent=True) or {}
    ids_raw = data.get("ids")
    if isinstance(ids_raw, str):
        ids = [ids_raw]
    elif isinstance(ids_raw, list):
        ids = [str(value) for value in ids_raw if value]
    else:
        ids = []

    if not ids:
        return jsonify({"error": "Missing notification ids"}), 400

    updated = 0
    with NOTIFICATION_LOCK:
        notifications = _load_notifications()
        id_set = set(ids)
        for entry in notifications:
            entry_id = entry.get("id")
            if entry_id in id_set and not entry.get("read"):
                entry["read"] = True
                updated += 1
        _write_notifications(notifications)

    return jsonify({"status": "ok", "updated": updated})


@app.post("/api/scheduler")
def scheduler_create() -> ResponseReturnValue:
    data = request.get_json(force=True, silent=True) or {}
    vendor_id = str(data.get("vendor_id") or "").strip()
    frequency = str(data.get("frequency") or "").strip().lower()
    start_date = data.get("start_date") or data.get("start_at")

    if not vendor_id:
        return jsonify({"error": "vendor_id is required"}), 400
    if frequency not in SCHEDULE_FREQUENCIES:
        return jsonify({"error": "frequency must be daily, weekly, or monthly"}), 400

    start_at = _parse_start_date(start_date)
    created_at = datetime.now(timezone.utc)
    next_run = _determine_next_run(start_at, frequency, created_at)

    schedule = {
        "id": str(uuid.uuid4()),
        "vendor_id": vendor_id,
        "frequency": frequency,
        "start_at": _isoformat_utc(start_at),
        "created_at": _isoformat_utc(created_at),
        "next_run_at": _isoformat_utc(next_run) if next_run else None,
    }

    with SCHEDULE_LOCK:
        schedules = _load_schedules()
        schedules.append(schedule)
        _write_schedules(schedules)

    return jsonify({"schedule": schedule}), 201


@app.get("/api/scheduler")
def scheduler_list() -> ResponseReturnValue:
    with SCHEDULE_LOCK:
        schedules = _load_schedules()
    return jsonify({"schedules": schedules})


@app.delete("/api/scheduler/<schedule_id>")
def scheduler_delete(schedule_id: str) -> ResponseReturnValue:
    with SCHEDULE_LOCK:
        schedules = _load_schedules()
        remaining = [entry for entry in schedules if entry.get("id") != schedule_id]
        if len(remaining) == len(schedules):
            return jsonify({"error": "Schedule not found"}), 404
        _write_schedules(remaining)
    return "", 204


@app.post("/api/reports/generate")
def reports_generate() -> ResponseReturnValue:
    data = request.get_json(force=True, silent=True) or {}
    run_id = str(data.get("run_id") or "").strip()
    requested_format = str(data.get("format") or "").strip().lower()

    if not run_id:
        return jsonify({"error": "run_id is required"}), 400
    if requested_format not in {"pdf", "docx"}:
        return jsonify({"error": "format must be 'pdf' or 'docx'"}), 400

    snapshot = _resolve_run_snapshot(run_id)
    if not snapshot:
        return jsonify({"error": "Run not found"}), 404

    content = _compose_report_content(snapshot)
    reports_dir = _reports_dir()
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    filename = secure_filename(f"report_{run_id}_{timestamp}.{requested_format}")
    if not filename:
        filename = f"report_{timestamp}.{requested_format}"
    report_path = (reports_dir / filename).resolve()
    reports_root = reports_dir.resolve()
    if reports_root not in report_path.parents:
        report_path = reports_root / filename

    try:
        if requested_format == "pdf":
            _render_pdf(report_path, content)
        else:
            _render_docx(report_path, content)
    except RuntimeError as exc:
        LOGGER.exception("Failed to generate report for run %s", run_id)
        return jsonify({"error": str(exc)}), 500

    response_payload = {
        "status": "ok",
        "run_id": run_id,
        "format": requested_format,
        "report_id": report_path.stem,
        "generated_at": content.get("generated_at"),
        "url": f"/api/reports/download/{report_path.name}",
        "metrics": content.get("metrics"),
    }

    return jsonify(response_payload)


@app.get("/api/reports/download/<path:filename>")
def reports_download(filename: str) -> ResponseReturnValue:
    reports_dir = _reports_dir()
    safe_name = Path(filename).name
    if not safe_name:
        return jsonify({"error": "Report not found"}), 404

    reports_root = reports_dir.resolve()
    candidate_path = (reports_dir / safe_name).resolve()

    if not candidate_path.exists() or not candidate_path.is_file() or reports_root not in candidate_path.parents:
        return jsonify({"error": "Report not found"}), 404

    return send_file(candidate_path, as_attachment=True)


@app.post("/api/policy/check")
def policy_check() -> ResponseReturnValue:
    if "file" not in request.files:
        return jsonify({"error": "Missing 'file' upload field"}), 400

    upload = request.files["file"]
    if upload is None or not upload.filename or not upload.filename.strip():
        return jsonify({"error": "Policy document filename missing"}), 400

    controls = [value for value in request.form.getlist("controls") if value and value.strip()]
    policy_run_id = f"POLICY-{_new_run_id()}"
    policy_root = OUT_DIR / "policy_checks"
    policy_dir = policy_root / policy_run_id
    policy_dir.mkdir(parents=True, exist_ok=True)

    safe_name = secure_filename(upload.filename) or f"policy_{policy_run_id}.txt"
    saved_path = policy_dir / safe_name

    try:
        upload.save(saved_path)
    except OSError as exc:  # pragma: no cover - disk errors are rare in tests
        LOGGER.exception("Failed to persist uploaded policy document")
        return jsonify({"error": "Unable to save uploaded document"}), 500

    try:
        analysis = analyze_policy_document(saved_path, controls or None)
    except FileNotFoundError as exc:
        LOGGER.exception("Policy document missing during analysis")
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:  # pragma: no cover - defensive logging for unexpected errors
        LOGGER.exception("Policy analysis failed")
        return jsonify({"error": "Policy analysis failed"}), 500

    analysis["document_name"] = upload.filename
    analysis.setdefault("controls_evaluated", list(DEFAULT_CONTROL_IDS))
    response_payload = {"policy_run_id": policy_run_id, **analysis}
    return jsonify(response_payload)


@app.post("/api/nl_query")
def nl_query() -> ResponseReturnValue:
    data = request.get_json(force=True, silent=True) or {}
    text = (data.get("text") or "").strip()
    email = (data.get("email") or "").strip() or None
    if not text:
        return jsonify({"error": "Missing 'text' payload"}), 400

    parsed_intent = intent_parser.parse_intent(text, use_llm=False)
    plan = planner_agent.plan_tasks(parsed_intent)

    run_id = _new_run_id()
    run_queue: "queue.Queue[Dict[str, Any] | None]" = queue.Queue()
    with RUN_LOCK:
        RUN_QUEUES[run_id] = run_queue

    thread = threading.Thread(target=_worker, args=(run_id, parsed_intent, plan, email, text), daemon=True)
    thread.start()

    response = {
        "run_id": run_id,
        "stream_url": f"/api/stream/{run_id}",
    }
    if email:
        response["email"] = email

    return jsonify(response)


@app.get("/api/stream/<run_id>")
def stream(run_id: str) -> ResponseReturnValue:
    with RUN_LOCK:
        q = RUN_QUEUES.get(run_id)
    if q is None:
        return jsonify({"error": "Unknown run"}), 404

    def event_stream():
        while True:
            item = q.get()
            if item is None:
                yield "event: end\ndata: {}\n\n"
                break
            yield f"data: {json.dumps(item)}\n\n"

    return Response(stream_with_context(event_stream()), mimetype="text/event-stream")


@app.get("/api/download/<run_id>")
def download(run_id: str) -> ResponseReturnValue:
    result = RUN_RESULTS.get(run_id)
    if not result:
        return jsonify({"error": "Run not found"}), 404
    package_path = result.get("package_path")
    if not package_path or not Path(package_path).exists():
        return jsonify({"error": "Package not available"}), 404
    return send_file(package_path, as_attachment=True)


@app.post("/api/confirm_send")
def confirm_send() -> ResponseReturnValue:
    data = request.get_json(force=True, silent=True) or {}
    run_id = data.get("run_id")
    recipient = (data.get("email") or "").strip()
    if not run_id or not recipient:
        return jsonify({"error": "run_id and email are required"}), 400

    result = RUN_RESULTS.get(run_id)
    if not result:
        return jsonify({"error": "Run not found"}), 404

    package_path = result.get("package_path")
    if not package_path or not Path(package_path).exists():
        return jsonify({"error": "Package not available"}), 400

    manifest = result.get("manifest", {})
    notify = orchestrator.notify_agent
    send_result = notify.send_package(
        run_manifest=manifest,
        recipient_email=recipient,
        attachments=[package_path],
    )

    return jsonify({"status": "sent", "details": send_result})


@app.get("/api/runs/<run_id>/conversation")
def conversation(run_id: str) -> ResponseReturnValue:
    path = _conversation_path(run_id)
    if path.exists():
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            LOGGER.warning("Conversation log for %s is invalid JSON", run_id)
        else:
            return jsonify(data)

    result = RUN_RESULTS.get(run_id)
    if result and "conversation" in result:
        return jsonify({"run_id": run_id, "messages": result["conversation"]})

    return jsonify({"error": "Conversation not found", "run_id": run_id}), 404


if __name__ == "__main__":
    app.run(debug=True)
